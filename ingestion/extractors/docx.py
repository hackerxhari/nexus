"""
DOCX extractor — handles Word documents.
Extracts paragraphs, tables, and headers preserving document structure.
"""

import os
import zipfile
from typing import List

from docx import Document
from docx.oxml.ns import qn

from core.exceptions import ExtractionFailedError
from core.logger import get_logger
from ingestion.extractors.base import BaseExtractor, ExtractionResult

logger = get_logger(__name__)


class DOCXExtractor(BaseExtractor):

    @property
    def supported_extensions(self) -> List[str]:
        return ["docx", "doc"]

    def _extract(self, filepath: str) -> ExtractionResult:
        if not os.path.exists(filepath):
            raise ExtractionFailedError(
                os.path.basename(filepath),
                "File not found"
            )

        if not zipfile.is_zipfile(filepath):
            raise ExtractionFailedError(
                os.path.basename(filepath),
                "Not a valid DOCX file. If this is a .doc file, please convert to .docx."
            )

        warnings = []
        sections = []

        try:
            doc = Document(filepath)

            # Extract paragraphs preserving heading structure
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Mark headings for better chunking context
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    sections.append(f"\n## {text}\n")
                else:
                    sections.append(text)

            # Extract tables as structured text
            for table_idx, table in enumerate(doc.tables):
                try:
                    table_text = self._extract_table(table)
                    if table_text:
                        sections.append(f"\n[Table {table_idx + 1}]\n{table_text}\n")
                except Exception as e:
                    warnings.append(f"Table {table_idx + 1} failed: {str(e)}")

        except Exception as e:
            message = str(e)
            lower = message.lower()
            if "password" in lower or "encrypted" in lower:
                message = "File appears to be encrypted or password-protected."
            raise ExtractionFailedError(
                os.path.basename(filepath),
                f"DOCX parsing failed: {message}"
            )

        combined_text = "\n".join(sections)

        return ExtractionResult(
            text=combined_text,
            page_count=0,  # DOCX doesn't have reliable page count
            extraction_method="docx_parser",
            warnings=warnings,
            metadata={"table_count": len(doc.tables)}
        )

    def _extract_table(self, table) -> str:
        """Convert a DOCX table to readable text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Remove duplicate adjacent cells (merged cells repeat content)
            deduped = [cells[0]] + [
                cells[i] for i in range(1, len(cells))
                if cells[i] != cells[i - 1]
            ]
            row_text = " | ".join(c for c in deduped if c)
            if row_text:
                rows.append(row_text)
        return "\n".join(rows)